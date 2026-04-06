#!/usr/bin/env python3
"""Generate WattsOn Competitive Analysis & Market Research Word Document."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.color.rgb = RGBColor(0x0A, 0x0E, 0x1A)
    heading_style.font.name = 'Calibri'

doc.styles['Heading 1'].font.size = Pt(24)
doc.styles['Heading 2'].font.size = Pt(18)
doc.styles['Heading 3'].font.size = Pt(14)

CYAN = RGBColor(0x00, 0xD4, 0xFF)
DARK_BG = RGBColor(0x0A, 0x0E, 0x1A)
GREEN = RGBColor(0x00, 0xAA, 0x55)
RED = RGBColor(0xCC, 0x33, 0x33)

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)
        set_cell_shading(cell, "0A0E1A")

    # Data rows
    for r, row_data in enumerate(rows):
        for c, value in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
            if r % 2 == 0:
                set_cell_shading(cell, "F0F4F8")

    return table

def add_bullet(doc, text, bold_prefix=None, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 1.27)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)

# ══════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════

# Add some spacing
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("WattsOn")
run.font.size = Pt(48)
run.bold = True
run.font.color.rgb = CYAN

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Competitive Analysis & Market Research")
run.font.size = Pt(24)
run.font.color.rgb = DARK_BG

doc.add_paragraph()

tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tagline.add_run("Egypt's First AI-Powered EV Charging Aggregator")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.italic = True

doc.add_paragraph()
doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run(f"Prepared: {datetime.date.today().strftime('%B %Y')}")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

info2 = doc.add_paragraph()
info2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info2.add_run("Confidential")
run.font.size = Pt(12)
run.bold = True
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ══════════════════════════════════════════════════════════════════════════

doc.add_heading('Table of Contents', level=1)
toc_items = [
    "1. Executive Summary",
    "2. Egypt EV Market Overview",
    "3. MENA Regional Market",
    "4. Egypt Charging Infrastructure",
    "5. Competitive Landscape — Egypt & MENA",
    "6. Competitive Landscape — Global Aggregators",
    "7. Competitive Positioning Matrix",
    "8. WattsOn's Unique Differentiators",
    "9. Business Models & Monetization",
    "10. Market Sizing (TAM / SAM / SOM)",
    "11. SWOT Analysis",
    "12. Growth Strategy & Expansion Roadmap",
    "13. Strategic Risks & Mitigations",
    "14. Recommended Monetization Roadmap",
    "15. Sources & References",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════

doc.add_heading('1. Executive Summary', level=1)

doc.add_paragraph(
    "WattsOn is Egypt's first and only AI-powered EV charging aggregator — a single app that "
    "maps all 435 charging stations across 12 providers and 16 governorates. In a market where "
    "EV drivers must juggle multiple provider-specific apps to find a charger, WattsOn provides "
    "a unified, intelligent layer with features no competitor offers: a Claude AI copilot, "
    "proximity-locked community reporting, smart trip planning, and a 329-vehicle dashboard "
    "with AI battery health analysis."
)

doc.add_paragraph(
    "This document presents a comprehensive competitive analysis and market research for WattsOn, "
    "covering the Egypt and MENA EV markets, infrastructure landscape, competitive positioning "
    "against both local and global players, business model benchmarks, and a strategic growth roadmap."
)

doc.add_heading('Key Findings', level=2)

findings = [
    ("Market Inflection Point: ", "Egypt's EV fleet doubled in 2024 (2,938 new EVs), with the total market valued at USD 10.22B and projected to reach USD 20.08B by 2030 (12% CAGR). Government has allocated EGP 1.5B for infrastructure and imposed 0% customs duty on EV imports."),
    ("Zero Competition in Aggregation: ", "No other app in Egypt or MENA aggregates charging stations across providers. All existing apps (InfinityEV, Sha7en) are single-network only. WattsOn has an 18-24 month window before global players could localize."),
    ("AI Differentiation is Unique Globally: ", "No EV charging app worldwide — not PlugShare (3.5M users), not Zap-Map (2M downloads), not ChargePoint — offers an AI copilot. WattsOn's Claude-powered assistant is a category-first feature."),
    ("Revenue Model Validated by Global Peers: ", "Zap-Map raised GBP 9M at a GBP 26.3M valuation. PlugShare was acquired for $25M. Multiple proven monetization paths exist: subscriptions, transaction fees (5-20%), advertising, data licensing, and fleet B2B."),
    ("MENA Expansion Opportunity: ", "The Middle East EV market is projected to reach USD 52.24B by 2030 (38% CAGR), with UAE and Saudi Arabia investing billions in charging infrastructure. No dominant aggregator exists in the region."),
]

for bold, text in findings:
    add_bullet(doc, text, bold)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 2. EGYPT EV MARKET OVERVIEW
# ══════════════════════════════════════════════════════════════════════════

doc.add_heading('2. Egypt EV Market Overview', level=1)

doc.add_heading('2.1 Current Market Status', level=2)

doc.add_paragraph(
    "Egypt has emerged as Africa's leading EV market, surpassing South Africa and Morocco in 2024. "
    "The country registered approximately 6,150 EVs by December 2024, with 2,938 purchased in 2024 alone — "
    "effectively doubling the total fleet in a single year (164% annual growth). By mid-2024, there were "
    "7,213 licensed EVs, with 1,419 vehicles registered in Q1 2024 alone. Q1 2025 saw an additional 39.5% "
    "increase in EV registrations over the prior quarter."
)

add_styled_table(doc,
    ["Metric", "Value"],
    [
        ["Total EVs Registered (Dec 2024)", "~6,150"],
        ["2024 EV Purchases", "2,938 (164% YoY growth)"],
        ["EV Penetration (total fleet)", "< 0.1%"],
        ["Market Valuation (2024)", "USD 10.22 billion"],
        ["Projected Market (2030)", "USD 20.08 billion"],
        ["CAGR (2024-2030)", "12.03%"],
        ["Rank in Africa", "#1 (surpassed South Africa)"],
        ["Consumer EV Consideration", "30% of car buyers"],
    ]
)

doc.add_paragraph()

doc.add_heading('2.2 Government Policies & Incentives', level=2)

policies = [
    ("0% customs duties on EV imports ", "(in effect since 2018, reinforced January 2025)"),
    ("Progressive FOB deduction: ", "10% initial + 10% per year of manufacture age, capping at 50%"),
    ("Used EV imports allowed: ", "vehicles up to 3 years old (unusual for Egypt)"),
    ("14% VAT ", "applies, plus shipping and registration fees"),
    ("EGP 1.5 billion (~USD 95M) ", "allocated for EV infrastructure development"),
    ("Government target: ", "3,000+ public EV chargers by 2027"),
    ("Mandatory EV charging ", "in new residential and commercial developments"),
    ("Local manufacturing push: ", "Geely opened Egypt's first CKD factory in January 2025 (10,000 vehicles/year capacity)"),
]

for bold, text in policies:
    add_bullet(doc, text, bold)

doc.add_heading('2.3 Market Drivers', level=2)

doc.add_paragraph(
    "The economics of EV ownership in Egypt are becoming increasingly compelling. Petrol prices reached "
    "EGP 21/liter by March 2026 (up 53% in 18 months), while EV charging costs offer up to 60% savings:"
)

add_styled_table(doc,
    ["Cost Comparison", "Petrol Car", "EV (Home Charging)", "Savings"],
    [
        ["Cost per 100 km", "EGP 168", "EGP 18-27", "~85%"],
        ["Annual cost (15,000 km)", "EGP 25,200", "EGP 2,700-4,050", "~84%"],
        ["Per full charge/tank", "EGP 400-600", "EGP 50-90 (home)", "~85%"],
        ["DC public charging", "—", "EGP 150-250/session", "~60% vs petrol"],
    ]
)

doc.add_paragraph()

doc.add_heading('2.4 Market Barriers', level=2)

barriers = [
    ("Infrastructure gap: ", "Fewer than 500 public stations for 110+ million people"),
    ("Urban/rural imbalance: ", "57% of population lives in rural areas with minimal charging access"),
    ("Low charger output: ", "Most stations max at 48 kW, below European fast-charging standards"),
    ("Only 5 licensed operators: ", "Limited competition in charging provision"),
    ("Connector fragmentation: ", "Type 2 dominates at 77.5%; CCS2, CHAdeMO, and GB/T face limited support"),
    ("High upfront cost: ", "Even with 0% duty, 14% VAT and fees add up; limited EV financing options"),
    ("Currency volatility: ", "EGP devaluation increases import costs"),
    ("Range anxiety: ", "Sparse intercity charging remains a major concern"),
]

for bold, text in barriers:
    add_bullet(doc, text, bold)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 3. MENA REGIONAL MARKET
# ══════════════════════════════════════════════════════════════════════════

doc.add_heading('3. MENA Regional Market', level=1)

doc.add_paragraph(
    "The broader MENA region represents a massive growth opportunity. The Middle East EV market was valued "
    "at USD 7.57 billion in 2024, projected to reach USD 52.24 billion by 2030 at a CAGR of 38%. "
    "The GCC EV market alone is projected to reach USD 9.42 billion by 2029."
)

doc.add_heading('3.1 Country Comparison', level=2)

add_styled_table(doc,
    ["Metric", "UAE", "Saudi Arabia", "Egypt"],
    [
        ["2024 Market Revenue", "USD 2.22B", "Fastest growing", "USD 10.22B"],
        ["Projected 2030", "USD 16.31B", "Highest CAGR (32.1%)", "USD 20.08B"],
        ["CAGR", "39.4%", "32.1%", "12.03%"],
        ["EV Target", "50% by 2050", "30% of Riyadh by 2030", "3,000 chargers by 2027"],
        ["Key Investment", "AED 60M (UAEV JV)", "SAR 150B (PIF)", "EGP 1.5B (gov't)"],
    ]
)

doc.add_paragraph()

doc.add_heading('3.2 Regional Investment Flows', level=2)

investments = [
    ("Saudi PIF: ", "SAR 150 billion (USD 40B) toward EV infrastructure and manufacturing"),
    ("Saudi PIF (charging): ", "SAR 5.3 billion (USD 1.41B) for charging infrastructure through 2025"),
    ("UAE UAEV: ", "AED 60 million initial capital, targeting 1,000 stations by 2030"),
    ("Dubai Municipality: ", "AED 150 million for public space EV charging"),
    ("MENA VC activity: ", "USD 5 billion total startup funding in 2024; cleantech gaining momentum"),
    ("DFI funding: ", "EBRD invested USD 101.5M in Infinity Energy (Egypt)"),
]

for bold, text in investments:
    add_bullet(doc, text, bold)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 4. EGYPT CHARGING INFRASTRUCTURE
# ══════════════════════════════════════════════════════════════════════════

doc.add_heading('4. Egypt Charging Infrastructure', level=1)

doc.add_paragraph(
    "Egypt's charging infrastructure is growing rapidly but remains concentrated in major cities. "
    "Fewer than 500 functional charging stations exist nationally, with only 5 companies licensed to operate. "
    "WattsOn has mapped 435 stations — potentially the most comprehensive database in the country."
)

doc.add_heading('4.1 Key Providers', level=2)

add_styled_table(doc,
    ["Provider", "Stations/Points", "Coverage", "Notes"],
    [
        ["Infinity EV", "135 stations, 700+ points", "16 governorates", "Egypt's largest; targeting 1,000 points by end 2025"],
        ["Sha7en", "~50-100 est.", "Multiple cities", "Manages former Revolta stations; AMPECO-powered"],
        ["Revolta Egypt", "65 (legacy)", "7 governorates", "Ceased operations; stations taken over by others"],
        ["New Energy", "Growing", "Cairo focus", "Provides charging guides"],
        ["Elsewedy Plug", "Various", "Limited", "Infrastructure/hardware focused"],
    ]
)

doc.add_paragraph()

doc.add_heading('4.2 Major Infrastructure Partnerships', level=2)

doc.add_paragraph(
    "Hassan Allam Utilities + Infinity + Misr Petroleum + Gastec signed a shareholders agreement to deploy "
    "EV chargers at existing fuel stations — leveraging extensive gas station networks for high-traffic locations. "
    "Infinity alone has raised USD 101.5 million total, with EBRD as a key investor. "
    "This partnership targets 3,000 twin chargers along the Cairo-Alexandria corridor."
)

doc.add_heading('4.3 Coverage Gaps', level=2)

gaps = [
    ("Urban concentration: ", "Vast majority of stations in Cairo and Alexandria"),
    ("Rural deficit: ", "57% of population in rural areas with minimal/no access"),
    ("Upper Egypt: ", "Only EGP 60M dedicated investment — progress is slow"),
    ("Highway corridors: ", "Limited intercity infrastructure, major barrier to long-distance travel"),
    ("Charger output: ", "Most stations max at 48 kW — slower than international standards"),
]

for bold, text in gaps:
    add_bullet(doc, text, bold)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 5. COMPETITIVE LANDSCAPE — EGYPT & MENA
# ══════════════════════════════════════════════════════════════════════════

doc.add_heading('5. Competitive Landscape — Egypt & MENA', level=1)

# InfinityEV
doc.add_heading('5.1 InfinityEV (Egypt)', level=2)

doc.add_paragraph(
    "InfinityEV is the app for Infinity Energy's charging network — Egypt's largest EV charging operator "
    "and subsidiary of Infinity Energy, one of Egypt's biggest renewable energy companies."
)

add_styled_table(doc,
    ["Attribute", "Details"],
    [
        ["Type", "Single-network operator app"],
        ["Network Size", "240+ stations, 700+ charging points"],
        ["Key Features", "Real-time availability, start/stop charging, in-app wallet, Valu BNPL"],
        ["Business Model", "Revenue from electricity sales at owned stations"],
        ["Funding", "Part of Infinity Energy (USD 101.5M from EBRD)"],
        ["Strengths", "Largest network, deep corporate backing, fuel station rollout"],
        ["Weaknesses", "Single-network only, no aggregation, no AI, no trip planning, no community features"],
    ]
)

doc.add_paragraph()

# Sha7en
doc.add_heading('5.2 Sha7en (Egypt)', level=2)

add_styled_table(doc,
    ["Attribute", "Details"],
    [
        ["Type", "Single-network operator app (AMPECO-powered)"],
        ["Network Size", "~50-100 stations (incl. former Revolta)"],
        ["Key Features", "Station locator, in-app wallet, home charger integration, loyalty program"],
        ["Business Model", "Revenue from charging sessions"],
        ["Strengths", "Modern platform (AMPECO), loyalty program, home charging"],
        ["Weaknesses", "Single-network, smaller than Infinity, no aggregation/AI/trip planning"],
    ]
)

doc.add_paragraph()

# Shabik/CATEC
doc.add_heading('5.3 Shabik / CATEC (UAE, MENA)', level=2)

add_styled_table(doc,
    ["Attribute", "Details"],
    [
        ["Type", "Single-network operator (MENA's largest)"],
        ["Coverage", "UAE, Saudi Arabia, Jordan — NOT in Egypt"],
        ["Key Features", "15-min plug reservation, RFID card, real-time monitoring, 24/7 support"],
        ["Strengths", "Multi-country presence, reservation feature"],
        ["Weaknesses", "Not in Egypt, single-network, separate apps per country, no AI or community"],
    ]
)

doc.add_paragraph()

doc.add_heading('5.4 Key Insight: No Aggregator Exists', level=2)

doc.add_paragraph(
    "The critical finding is that no competitor in Egypt or MENA operates as a true charging aggregator. "
    "Every existing app — InfinityEV, Sha7en, Shabik — shows only their own network's stations. "
    "An Egyptian EV driver today must download and maintain multiple apps to find the nearest charger. "
    "WattsOn is the only app that solves this fragmentation problem by aggregating all 12 providers into a single view."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 6. COMPETITIVE LANDSCAPE — GLOBAL AGGREGATORS
# ══════════════════════════════════════════════════════════════════════════

doc.add_heading('6. Competitive Landscape — Global Aggregators', level=1)

doc.add_paragraph(
    "Globally, the EV charging app market has shifted decisively from single-network proprietary apps to "
    "cross-network aggregation platforms. The charging network roaming hub market is valued at USD 420 million "
    "in 2024, projected to hit USD 2.78 billion by 2033 (CAGR 23.2%)."
)

# PlugShare
doc.add_heading('6.1 PlugShare', level=2)

add_styled_table(doc,
    ["Attribute", "Details"],
    [
        ["Users", "3.5M+ registered (growing 40%+ annually)"],
        ["Stations", "800,000+ mapped globally"],
        ["Community", "5.8M+ reviews, 725K+ photos, 6.5M+ check-ins"],
        ["Business Model", "Freemium ($14.99/yr ad-free) + data licensing + API licensing to automakers"],
        ["Acquisition", "Acquired by EVgo for $25M (2021)"],
        ["Strengths", "Largest crowd-sourced EV database, massive network effects, B2B data"],
        ["Weaknesses", "Minimal Egypt/MENA coverage, no AI, no Arabic, basic trip planning"],
    ]
)

doc.add_paragraph()

# Zap-Map
doc.add_heading('6.2 Zap-Map (UK)', level=2)

add_styled_table(doc,
    ["Attribute", "Details"],
    [
        ["Users", "2M+ downloads, ~70% UK EV driver penetration"],
        ["Coverage", "87,000+ charge devices (~85% of UK public infrastructure)"],
        ["Business Model", "Freemium + Zap-Pay transaction fees + B2B data"],
        ["Funding", "GBP 9M Series A (2022), valued at GBP 26.3M"],
        ["Strengths", "UK market dominance, cross-network payment, CarPlay/Android Auto"],
        ["Weaknesses", "UK-only, no AI, no vehicle dashboard, no Arabic"],
    ]
)

doc.add_paragraph()

# ABRP
doc.add_heading('6.3 A Better Route Planner (ABRP)', level=2)

add_styled_table(doc,
    ["Attribute", "Details"],
    [
        ["Users", "1.5M+ downloads, 100M+ trips planned"],
        ["Focus", "Best-in-class EV trip planning with vehicle-specific models"],
        ["Business Model", "Freemium (premium for live data, CarPlay)"],
        ["Acquisition", "Acquired by Rivian (June 2023)"],
        ["Strengths", "Gold standard for trip planning, detailed customization"],
        ["Weaknesses", "Narrow focus, no community reporting, no AI, no Egypt station data"],
    ]
)

doc.add_paragraph()

# Others summary
doc.add_heading('6.4 Other Global Players', level=2)

add_styled_table(doc,
    ["Player", "Type", "Users/Size", "Key Differentiator", "Egypt Presence"],
    [
        ["ChargePoint", "Network + App", "1M+ monthly drivers, 375K ports", "Largest US network, AI platform (B2B)", "None"],
        ["Chargeway", "Aggregator", "Niche", "Color-coded simplification, real-time pricing", "None"],
        ["Electromaps", "Aggregator", "860K+ users", "European coverage, B2B platform", "None"],
        ["Shell Recharge", "Oil Major App", "250K+ roaming points", "Fuel + EV integration, brand trust", "None"],
        ["bp pulse", "Oil Major App", "41K+ charging bays", "Airport charging hubs, subscription model", "None"],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 7. COMPETITIVE POSITIONING MATRIX
# ══════════════════════════════════════════════════════════════════════════

doc.add_heading('7. Competitive Positioning Matrix', level=1)

doc.add_paragraph(
    "The following matrix compares WattsOn against all major competitors across key capabilities. "
    "A checkmark indicates the feature is present; a dash indicates absence."
)

# Use Y/N for Word compatibility
add_styled_table(doc,
    ["Feature", "WattsOn", "InfinityEV", "Sha7en", "PlugShare", "Zap-Map", "ABRP"],
    [
        ["Multi-provider aggregation", "YES", "No", "No", "YES", "YES", "Partial"],
        ["AI copilot / assistant", "YES", "No", "No", "No", "No", "No"],
        ["Community status reporting", "YES", "No", "No", "YES", "YES", "No"],
        ["Proximity-locked reports", "YES", "No", "No", "No", "No", "No"],
        ["Trip planning", "YES", "No", "No", "Basic", "Basic", "BEST"],
        ["Vehicle dashboard", "YES", "No", "No", "No", "No", "Models"],
        ["AI battery health", "YES", "No", "No", "No", "No", "No"],
        ["Arabic + English", "YES", "Likely", "Likely", "No", "No", "No"],
        ["Egypt coverage", "435 stn", "240+", "~50-100", "Minimal", "None", "Minimal"],
        ["In-app payment", "No", "YES", "YES", "YES", "YES", "No"],
        ["Start/stop charging", "No", "YES", "YES", "Partial", "Partial", "No"],
        ["EV news feed", "YES", "No", "No", "No", "No", "No"],
        ["Crowdsource stations", "YES", "No", "No", "YES", "YES", "No"],
        ["User base", "Early", "Moderate", "Small", "3.5M+", "2M+", "1.5M+"],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 8. WATTSON'S UNIQUE DIFFERENTIATORS
# ══════════════════════════════════════════════════════════════════════════

doc.add_heading("8. WattsOn's Unique Differentiators", level=1)

differentiators = [
    ("Only True Aggregator in Egypt",
     "No other app in the Egyptian market aggregates stations across all 12 providers. InfinityEV only shows Infinity stations. Sha7en only shows Sha7en stations. WattsOn is the only app where an Egyptian EV driver can see all 435 stations in one place. This is the same value proposition that made PlugShare dominant in the US and Zap-Map dominant in the UK."),
    ("AI Copilot — A Global First",
     "The Claude-powered AI assistant is genuinely unique across the entire competitive landscape. No competitor — local or global — offers an AI copilot that understands real station data, can answer charging questions in context, and provides personalized assistance. ChargePoint has AI in their B2B platform, but nothing consumer-facing."),
    ("Proximity-Locked Community Reporting",
     "While PlugShare and Zap-Map have community reporting, WattsOn's 100-meter proximity lock ensures reports come from people physically at the station. This is a data quality innovation that no global player has implemented, creating a higher-trust dataset."),
    ("AI Battery Health Dashboard",
     "No competitor offers AI-powered battery health analysis integrated with a charging aggregator. The combination of a 329-vehicle database with personalized battery diagnostics is a novel feature category."),
    ("Arabic-First Bilingual Experience",
     "None of the global aggregators support Arabic. With 160+ translation keys across Arabic and English, WattsOn serves a market where Arabic is essential for mass adoption."),
    ("Integrated EV News Magazine",
     "Live RSS feeds from Electrek, InsideEVs, CleanTechnica, The Driven, and Green Car Reports create a content engagement loop that keeps users returning even when not actively charging. No competitor offers this."),
    ("Trip Planner That Works in Egypt",
     "While ABRP is superior in pure trip planning globally, it has virtually no Egypt station data. WattsOn's corridor-search trip planner is the only one that actually works for Egyptian routes with real station coverage."),
]

for title, desc in differentiators:
    doc.add_heading(title, level=3)
    doc.add_paragraph(desc)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 9. BUSINESS MODELS & MONETIZATION
# ══════════════════════════════════════════════════════════════════════════

doc.add_heading('9. Business Models & Monetization', level=1)

doc.add_paragraph(
    "Analysis of global EV charging aggregators reveals multiple proven revenue models. "
    "The industry average platform take-rate is approximately 19% for charging platforms, "
    "with rates ranging from 5-20% depending on market maturity."
)

doc.add_heading('9.1 Revenue Models', level=2)

models = [
    ("Transaction Fees (5-20%): ", "Commission on each charging session initiated through the app. The dominant model for aggregators with payment integration. Industry average take-rate: ~19%."),
    ("Subscription Tiers: ", "Freemium model with premium features. Benchmarks: Zap-Map GBP 34.99/yr, ABRP ~$60/yr, PlugShare $14.99/yr ad-free."),
    ("Advertising & Sponsored Listings: ", "Charging networks pay for promoted placement; EV brands and dealers target a highly qualified audience. MENA CPM rates: $5-15."),
    ("Data Licensing: ", "Anonymized charging patterns, demand heatmaps, and usage analytics sold to CPOs, utilities, government planners, and investors. Contracts typically $10K-$100K/year."),
    ("Fleet / B2B SaaS: ", "Fleet management dashboards for ride-hailing (Uber, Careem) and delivery companies transitioning to EVs. Pricing: $500-5,000/month per fleet."),
    ("Referral & Affiliate: ", "Commissions on EV purchases, insurance, home charger installations, and charging network sign-ups."),
]

for bold, text in models:
    add_bullet(doc, text, bold)

doc.add_heading('9.2 Pricing Benchmarks (PPP-Adjusted for Egypt)', level=2)

add_styled_table(doc,
    ["Tier", "Features", "Price (EGP)", "Price (USD equiv.)"],
    [
        ["WattsOn Free", "Full map, station info, community reports, basic AI", "Free", "Free"],
        ["WattsOn Plus", "Ad-free, advanced filters, priority AI, CarPlay", "EGP 29/mo or 249/yr", "~$0.60/mo or $5/yr"],
        ["WattsOn Pro", "Fleet features, trip planner pro, data exports, API access", "EGP 59/mo or 499/yr", "~$1.20/mo or $10/yr"],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 10. MARKET SIZING
# ══════════════════════════════════════════════════════════════════════════

doc.add_heading('10. Market Sizing (TAM / SAM / SOM)', level=1)

doc.add_heading('10.1 Total Addressable Market (TAM)', level=2)

doc.add_paragraph(
    "Based on an estimated 30,000 EVs on Egyptian roads by end of 2026, averaging 3 charges/week "
    "at EGP 150/session, the gross annual charging volume is approximately EGP 702 million (~USD 14.4M). "
    "At a 10% aggregator take rate, the TAM for a charging aggregator is ~USD 1.44M in transaction revenue alone."
)

doc.add_heading('10.2 Serviceable Addressable Market (SAM)', level=2)

doc.add_paragraph(
    "Adjusting for public charging only (40% of sessions), smart-capable stations (70%), "
    "and WattsOn's geographic coverage (80%), the SAM is approximately $3.2M in gross charging volume. "
    "Adding subscription, advertising, and data revenue, the total SAM reaches $4-6 million by 2027."
)

doc.add_heading('10.3 Serviceable Obtainable Market (SOM)', level=2)

add_styled_table(doc,
    ["Year", "MAU", "Premium Subs", "Revenue Estimate", "Revenue Sources"],
    [
        ["Year 1 (2026-27)", "5,000", "500", "$30K-60K", "Subscriptions + ads"],
        ["Year 2 (2027-28)", "15,000", "2,000", "$150K-300K", "+ Transaction fees"],
        ["Year 3 (2028-29)", "40,000", "5,000", "$500K-1M", "+ Fleet B2B + data licensing"],
    ]
)

doc.add_paragraph()
doc.add_paragraph(
    "Key assumptions: EV fleet doubles every 18-24 months; WattsOn captures 30-50% of Egypt's EV driver "
    "app market (first-mover advantage); payment integration launches in Year 2; 3-5 fleet contracts by Year 3."
)

doc.add_heading('10.4 MENA Expansion TAM', level=2)

doc.add_paragraph(
    "The MEA EV market is expected to reach USD 20.39 billion by 2031 (CAGR 32.15%). "
    "The MEA EV charging infrastructure market is at USD 165 million in 2025, growing at 16.4% CAGR. "
    "The total MENA aggregator TAM is estimated at $50-100M by 2030."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 11. SWOT ANALYSIS
# ══════════════════════════════════════════════════════════════════════════

doc.add_heading('11. SWOT Analysis', level=1)

doc.add_heading('Strengths', level=2)
strengths = [
    ("First-mover monopoly: ", "Zero direct competitors as an EV charging aggregator in Egypt"),
    ("Comprehensive coverage: ", "435 stations across 16 governorates, 12 providers"),
    ("AI differentiation: ", "Claude-powered copilot — no competitor offers this globally"),
    ("Community moat: ", "Proximity-locked reporting creates compounding data advantage"),
    ("Vehicle database: ", "329 models from 88 brands for personalized recommendations"),
    ("Bilingual: ", "160+ Arabic/English translation keys — critical for Egyptian market"),
    ("Modern tech stack: ", "React Native + Supabase + Vercel — fast iteration, cross-platform"),
    ("Live product: ", "Deployed at wattson-ev.vercel.app — not vaporware"),
]
for bold, text in strengths:
    add_bullet(doc, text, bold)

doc.add_heading('Weaknesses', level=2)
weaknesses = [
    ("Pre-revenue: ", "No monetization implemented yet"),
    ("Small team: ", "Solo founder with limited development capacity"),
    ("No payment integration: ", "Cannot process charging payments through the app"),
    ("Limited marketing budget: ", "Organic growth only, no paid acquisition"),
    ("No formal CPO agreements: ", "Station data relationships need to be formalized"),
    ("Web-only distribution: ", "Not yet in App Store or Play Store (Expo enables quick deployment)"),
]
for bold, text in weaknesses:
    add_bullet(doc, text, bold)

doc.add_heading('Opportunities', level=2)
opportunities = [
    ("Egypt EV boom: ", "15,000 EV sales expected in 2026, doubling year-over-year"),
    ("Government push: ", "EGP 1.5B allocated, 0% import duty, mandatory charging in new buildings"),
    ("Infinity expansion: ", "6,000 planned charging points at 3,000 stations"),
    ("Zero competitor aggregators: ", "18-24 month window in Egypt and MENA"),
    ("Fleet electrification: ", "Uber, Careem, delivery fleets beginning EV transition"),
    ("MENA expansion: ", "$5B+ market in 2026, growing at 32% CAGR"),
    ("Tourism corridors: ", "Red Sea, Sinai, Luxor routes need charging navigation"),
    ("Data monetization: ", "Unique dataset on Egyptian EV charging behavior"),
]
for bold, text in opportunities:
    add_bullet(doc, text, bold)

doc.add_heading('Threats', level=2)

add_styled_table(doc,
    ["Threat", "Severity", "Likelihood", "Mitigation"],
    [
        ["CPOs build own apps", "High", "High", "Cross-network value no single CPO can match"],
        ["Google/Apple Maps add charging", "Medium", "High", "Deeper features (community, AI, real-time)"],
        ["Global aggregator enters Egypt", "Medium", "Medium", "Local data moat, Arabic-first, community"],
        ["Slow EV adoption", "High", "Medium", "Diversify to PHEV, fleet B2B revenue"],
        ["Currency volatility (EGP)", "Medium", "High", "Price in EGP, USD from data/B2B"],
        ["Regulatory changes", "Low", "Low", "Position as infrastructure partner"],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 12. GROWTH STRATEGY & EXPANSION ROADMAP
# ══════════════════════════════════════════════════════════════════════════

doc.add_heading('12. Growth Strategy & Expansion Roadmap', level=1)

doc.add_heading('12.1 Network Effects Playbook', level=2)

effects = [
    ("Supply-side first: ", "Aggregate every charging station. Data completeness is the moat — WattsOn already has 435."),
    ("Community contributions: ", "User-reported status creates fresh, real-time data no single CPO can match."),
    ("Cross-side network effects: ", "More drivers reporting = more accurate data = more drivers using = more CPOs wanting visibility."),
    ("Data flywheel: ", "Usage data improves AI recommendations, which improves UX, which drives more usage."),
]
for bold, text in effects:
    add_bullet(doc, text, bold)

doc.add_heading('12.2 First-Mover Window', level=2)

doc.add_paragraph(
    "WattsOn's position mirrors PlugShare circa 2012-2014 in the US. PlugShare launched when the US had "
    "~5,000 public chargers and was later acquired for $25M. Egypt has ~1,500 charging points growing to 3,000+, "
    "with zero aggregator competitors. The window is 18-24 months before either a global player localizes "
    "or a local competitor emerges. Brand recognition, community data, and station relationships built now "
    "will be extremely difficult to replicate."
)

doc.add_heading('12.3 Priority Partnerships', level=2)

immediate = [
    ("Infinity: ", "Data sharing agreement, featured listing (700+ charging points)"),
    ("ADNOC: ", "Early partner for payment integration"),
    ("Egyptian Electricity Holding Company: ", "Government data partnership"),
    ("BYD Egypt / MG Egypt: ", "Pre-install WattsOn on new EV deliveries (captive audience)"),
]
for bold, text in immediate:
    add_bullet(doc, text, bold)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Growth Partnerships (Year 2-3):")
run.bold = True

growth = [
    ("Uber Egypt / Careem: ", "Fleet integration, driver onboarding"),
    ("Egyptian banks (CIB, QNB ALAHLI): ", "Co-branded EV charging cards"),
    ("Mall operators: ", "Sponsored charging navigation at City Centre, Mall of Egypt"),
]
for bold, text in growth:
    add_bullet(doc, text, bold)

doc.add_heading('12.4 Regional Expansion Roadmap', level=2)

add_styled_table(doc,
    ["Phase", "Timeline", "Market", "Strategy"],
    [
        ["Phase 1", "Now - 2027", "Egypt", "Dominate home market, 435+ stations, community moat"],
        ["Phase 2", "2027 - 2028", "UAE + Saudi Arabia", "Largest MENA EV markets, high willingness to pay"],
        ["Phase 3", "2028 - 2029", "Morocco, Jordan, Kenya", "Emerging EV markets, limited competition"],
        ["Phase 4", "2029 - 2030", "Pan-MENA + East Africa", "Regional aggregator, cross-border trip planning"],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 13. STRATEGIC RISKS & MITIGATIONS
# ══════════════════════════════════════════════════════════════════════════

doc.add_heading('13. Strategic Risks & Mitigations', level=1)

risks = [
    ("No In-App Payment or Session Control",
     "Every network-operator competitor lets users start/stop charging and pay through their app. "
     "WattsOn currently cannot initiate sessions. This is the single biggest functional gap. "
     "Mitigation: Prioritize payment integration with top CPOs (Infinity first) in 2027."),
    ("User Base Scale",
     "PlugShare has 3.5M+ users; WattsOn is early stage. Egypt's EV market (~6,000-10,000 EVs) limits "
     "near-term scale. Mitigation: The land-grab opportunity is open — focus on capturing EV driver "
     "mindshare before the market scales."),
    ("Infinity's Dominance Risk",
     "Infinity is deploying 3,000 twin chargers. If they achieve 80%+ market share, drivers may default "
     "to their app. Mitigation: WattsOn's aggregation value increases with market fragmentation, but "
     "even with one dominant provider, cross-network search + AI + community features add value."),
    ("Software-Only Model",
     "WattsOn has no hardware or energy revenue. Mitigation: Build the business around data, advertising, "
     "premium features, and B2B services — proven models at PlugShare and Zap-Map."),
]

for title, desc in risks:
    doc.add_heading(title, level=3)
    doc.add_paragraph(desc)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 14. RECOMMENDED MONETIZATION ROADMAP
# ══════════════════════════════════════════════════════════════════════════

doc.add_heading('14. Recommended Monetization Roadmap', level=1)

doc.add_heading('Phase 1: Foundation (Q2-Q4 2026)', level=2)
phase1 = [
    "Launch freemium model: Free tier + WattsOn Plus (EGP 29/mo)",
    "Implement in-app advertising (CPO promoted listings, EV dealer ads)",
    "Target: $2,000-5,000/month revenue",
    "Cost: Near-zero marginal cost on existing infrastructure",
]
for item in phase1:
    add_bullet(doc, item)

doc.add_heading('Phase 2: Transaction Revenue (2027)', level=2)
phase2 = [
    "Integrate payment processing with top CPOs (Infinity, Sha7en)",
    "Launch 'Charge with WattsOn' — 5-8% commission per session",
    "Introduce fleet dashboard (B2B SaaS, EGP 2,000-5,000/month per fleet)",
    "Target: $10,000-25,000/month revenue",
]
for item in phase2:
    add_bullet(doc, item)

doc.add_heading('Phase 3: Scale (2028-2029)', level=2)
phase3 = [
    "Data licensing contracts with CPOs, government, and investors",
    "MENA expansion (UAE, Saudi Arabia)",
    "White-label partnerships with automakers for in-car integration",
    "Target: $50,000-100,000/month revenue",
]
for item in phase3:
    add_bullet(doc, item)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 15. SOURCES & REFERENCES
# ══════════════════════════════════════════════════════════════════════════

doc.add_heading('15. Sources & References', level=1)

sources = [
    "TechSci Research — Egypt EV Market Report",
    "Statista — Electric Vehicles Egypt Outlook",
    "EV24 Africa — Electric Cars in Egypt 2026",
    "EV24 Africa — EV Incentives & Charging Infrastructure in Egypt",
    "EV24 Africa — Electric vs Petrol Cars in Egypt: Real Ownership Costs",
    "YoCharge — State of Electric Mobility in Egypt 2024",
    "New Energy — EV Charging Stations in Egypt 2025 Guide",
    "EVCandi — Egypt Public & Private Sectors Expand EV Charging",
    "MDPI — EV Adoption in Egypt: Feasibility, Challenges, Policy Directions",
    "Mordor Intelligence — MEA EV Market Report",
    "Grand View Research — MEA EV Market / EV Charging Infrastructure Market",
    "Inc Arabia — Surge of EVs in MENA Region",
    "GlobeNewsWire — MEA EV Market Worth $5B in 2026",
    "Precedence Research — EV Charging Infrastructure Market",
    "MarketIntelo — Charging Network Roaming Hub Market 2033",
    "GreenFlux — 2025: The Year OCPI Unified EV Charging",
    "Virta — OCPI Protocol Explained",
    "PlugShare / EVgo — Company About, 6.5M Check-Ins Milestone",
    "Zap-Map — UK EV Charging Statistics 2025, Plans Compared",
    "ChargePoint — Network Data 2026",
    "ABRP / Rivian — Acquisition Announcement",
    "Chargeway — 2.0 Launch, Real-Time Pricing",
    "Electromaps / Wallbox — Acquisition, Platform Overview",
    "Shell Recharge — Network Overview",
    "bp pulse — Network Overview",
    "InfinityEV — Google Play, infinityevcharge.com",
    "Sha7en — sha7en.co, Enterprise Press",
    "Shabik / CATEC — catec.ae",
    "EBRD — $40M Investment in Infinity",
    "Enterprise — EV Tariffs, Charging Tariff Hikes",
    "Hassan Allam — EV Charging Shareholders Agreement",
    "Geely — Egypt CKD Plant Launch",
    "PwC Strategy& — EV Charging Market Outlook 2025",
    "Driivz — 2026 EV Charging Industry Predictions",
    "Solidstudio — EV Charging Station Profit Margins, eRoaming Hubs",
    "AMPECO — Transaction Costs, White-Label Software",
    "6W Research — Egypt EV Market 2025-2031",
    "MAGNiTT — Q1 2025 MENA Venture Investment Report",
]

for source in sources:
    add_bullet(doc, source)

# ── Save ────────────────────────────────────────────────────────────────

output_path = "/Users/bistrocloud/Documents/EV Charging Aggregator/docs/WattsOn_Competitive_Analysis_Market_Research.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")
