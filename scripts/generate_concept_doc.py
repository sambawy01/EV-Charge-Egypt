"""
Generate the WattsOn Concept Document (DOCX) — investor/stakeholder-ready.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

BRAND_DARK  = RGBColor(0x0F, 0x17, 0x2A)   # deep navy
BRAND_GREEN = RGBColor(0x00, 0xC9, 0x6F)   # electric green
BRAND_GREY  = RGBColor(0x4A, 0x4A, 0x5A)   # body grey
BRAND_LIGHT = RGBColor(0x7A, 0x7A, 0x8A)   # lighter grey
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)

FONT_BODY = "Calibri"
FONT_HEAD = "Calibri"


def set_cell_shading(cell, color_hex: str):
    """Apply background shading to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_paragraph_spacing(paragraph, before=0, after=0, line=None):
    pPr = paragraph._p.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = parse_xml(f'<w:spacing {nsdecls("w")}/>')
        pPr.append(spacing)
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    if line:
        spacing.set(qn("w:line"), str(line))
        spacing.set(qn("w:lineRule"), "auto")


def styled_run(paragraph, text, bold=False, italic=False, size=11, color=BRAND_GREY, font=FONT_BODY):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = font
    return run


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = BRAND_DARK
        run.font.name = FONT_HEAD
    return h


def add_body(doc, text, bold_phrases=None, size=11):
    """Add a paragraph, optionally bolding specific phrases."""
    p = doc.add_paragraph()
    if bold_phrases:
        remaining = text
        for phrase in sorted(bold_phrases, key=lambda x: text.find(x)):
            idx = remaining.find(phrase)
            if idx == -1:
                continue
            if idx > 0:
                styled_run(p, remaining[:idx], size=size)
            styled_run(p, phrase, bold=True, size=size)
            remaining = remaining[idx + len(phrase):]
        if remaining:
            styled_run(p, remaining, size=size)
    else:
        styled_run(p, text, size=size)
    return p


def add_bullet(doc, text, bold_prefix=None, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    if bold_prefix:
        styled_run(p, bold_prefix, bold=True)
        styled_run(p, text)
    else:
        styled_run(p, text)
    return p


def add_page_break(doc):
    doc.add_page_break()


def make_metric_table(doc, metrics):
    """Create a styled 2-column metrics table."""
    table = doc.add_table(rows=len(metrics), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, (value, label) in enumerate(metrics):
        c0 = table.cell(i, 0)
        c1 = table.cell(i, 1)
        c0.width = Cm(4)
        c1.width = Cm(12)
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        styled_run(p0, value, bold=True, size=16, color=BRAND_GREEN)
        p1 = c1.paragraphs[0]
        styled_run(p1, label, size=11)
        set_cell_shading(c0, "0F172A")
        set_cell_shading(c1, "F8F9FA")
    return table


def set_narrow_margins(section):
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)


# ---------------------------------------------------------------------------
# Build Document
# ---------------------------------------------------------------------------

doc = Document()

# -- Default style tweaks --
style = doc.styles["Normal"]
style.font.name = FONT_BODY
style.font.size = Pt(11)
style.font.color.rgb = BRAND_GREY

for level in range(1, 5):
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = FONT_HEAD
    hs.font.color.rgb = BRAND_DARK

section = doc.sections[0]
set_narrow_margins(section)

# ============================================================
# COVER PAGE
# ============================================================

# Add spacing before title
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
styled_run(p, "WattsOn", bold=True, size=42, color=BRAND_GREEN, font=FONT_HEAD)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
styled_run(p, "Egypt's Intelligent EV Charging Platform", size=18, color=BRAND_DARK)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
styled_run(p, "Find. Charge. Drive Smart.", bold=True, italic=True, size=14, color=BRAND_LIGHT)

# Spacer
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
styled_run(p, "Concept Document", size=14, color=BRAND_DARK)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
styled_run(p, "March 2026", size=12, color=BRAND_GREY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
styled_run(p, "CONFIDENTIAL", bold=True, size=10, color=RGBColor(0xCC, 0x00, 0x00))

add_page_break(doc)

# ============================================================
# TABLE OF CONTENTS
# ============================================================

add_heading_styled(doc, "Table of Contents", level=1)

toc_items = [
    ("1.", "Executive Summary"),
    ("2.", "The Problem"),
    ("3.", "The Solution: WattsOn"),
    ("4.", "Key Features"),
    ("", "    4.1  Interactive Charging Map"),
    ("", "    4.2  Community-Driven Station Status"),
    ("", "    4.3  AI Copilot (Claude-Powered)"),
    ("", "    4.4  AI Trip Planner"),
    ("", "    4.5  Vehicle Dashboard"),
    ("", "    4.6  EV News Magazine"),
    ("", "    4.7  Submit a Station"),
    ("", "    4.8  Proximity Intelligence"),
    ("5.", "Unique Value Proposition"),
    ("6.", "Technology Stack"),
    ("7.", "Market Opportunity"),
    ("8.", "Roadmap"),
    ("9.", "Team & Contact"),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    if num:
        styled_run(p, f"{num}  ", bold=True, size=12, color=BRAND_DARK)
        styled_run(p, title, size=12, color=BRAND_DARK)
    else:
        styled_run(p, title, size=11, color=BRAND_GREY)

add_page_break(doc)

# ============================================================
# 1. EXECUTIVE SUMMARY
# ============================================================

add_heading_styled(doc, "1. Executive Summary", level=1)

add_body(doc,
    "WattsOn is Egypt's first AI-powered EV charging aggregator -- a single platform that maps "
    "every charging station across the country, provides real-time community-driven availability "
    "data, and uses artificial intelligence to optimize every aspect of the EV charging experience.",
    bold_phrases=["Egypt's first AI-powered EV charging aggregator"])

add_body(doc,
    "In a market with 300+ scattered charging points across 12+ providers with no unified "
    "visibility, WattsOn becomes the essential companion for every EV driver in Egypt.",
    bold_phrases=["essential companion for every EV driver in Egypt"])

doc.add_paragraph()  # spacer

# Metrics table
add_heading_styled(doc, "Key Metrics at a Glance", level=2)

metrics = [
    ("435+",  "Charging stations mapped across Egypt"),
    ("12+",   "Charging providers aggregated (Infinity EV, Elsewedy Plug, Sha7en, IKARUS, Revolta, KarmCharge, Electra, and more)"),
    ("16",    "Governorates covered"),
    ("140+",  "EV models supported across 41 brands"),
    ("Real-time", "Crowdsourced station availability -- community-driven data"),
    ("AI",    "Powered trip planning, battery optimization, and intelligent copilot"),
]
make_metric_table(doc, metrics)

add_page_break(doc)

# ============================================================
# 2. THE PROBLEM
# ============================================================

add_heading_styled(doc, "2. The Problem", level=1)

add_body(doc,
    "Egypt's electric vehicle market is at an inflection point. Government targets aim for "
    "100,000 EVs on the road by 2030, and major international brands -- BYD, Tesla, BMW, "
    "Hyundai, Mercedes -- are rapidly entering the market. Yet the charging infrastructure "
    "experience remains deeply fragmented.",
    bold_phrases=["100,000 EVs on the road by 2030", "deeply fragmented"])

problems = [
    ("Fragmented Infrastructure: ",
     "Charging stations are scattered across 12+ providers, each with their own app, their own map, and their own payment system. There is no single source of truth."),
    ("Range Anxiety: ",
     "EV drivers live with the constant fear of running out of charge with no nearby station. This is the #1 barrier to EV adoption in Egypt."),
    ("No Real-Time Data: ",
     "Drivers arrive at stations only to find them busy, broken, or offline. No app in Egypt provides real-time availability information."),
    ("No Intelligent Planning: ",
     "Long-distance trips require manual research -- which stations are on the route? Will they be available? How long will charging take? There are no answers."),
    ("No Community Feedback: ",
     "Station reliability, cleanliness, and condition are unknown. Drivers cannot share or access experiences from other EV owners."),
    ("Information Silos: ",
     "Vehicle-specific data like battery health, optimal charging patterns, and consumption analytics are not available to drivers in any meaningful way."),
]

for prefix, text in problems:
    add_bullet(doc, text, bold_prefix=prefix)

add_page_break(doc)

# ============================================================
# 3. THE SOLUTION
# ============================================================

add_heading_styled(doc, "3. The Solution: WattsOn", level=1)

add_body(doc,
    "WattsOn solves every pain point with a unified, intelligent platform that puts the "
    "EV driver at the center of the experience.")

solutions = [
    ("One Map, All Stations: ",
     "Every charging station in Egypt from every provider on one dark-themed, premium map."),
    ("Community-Driven Availability: ",
     "Users within 100m of a station report real-time status (Available, Busy, Out of Service). No other app in Egypt does this."),
    ("AI Copilot: ",
     "Claude-powered assistant that knows your vehicle, nearby stations, prices, and driving patterns. Ask anything in natural language."),
    ("Smart Trip Planner: ",
     "Plan any route in Egypt with optimized charging stops, real drive times via Google Maps, and nearby attractions at each stop."),
    ("Vehicle Intelligence: ",
     "AI-powered battery health analysis, consumption tracking, degradation estimates, and personalized optimization tips."),
]

for prefix, text in solutions:
    add_bullet(doc, text, bold_prefix=prefix)

doc.add_paragraph()

add_body(doc,
    "WattsOn is not just another map. It is an intelligent operating system for the EV "
    "driver's life -- from the moment they consider a trip, through the drive, at the "
    "charging station, and back home.",
    bold_phrases=["intelligent operating system for the EV driver's life"])

add_page_break(doc)

# ============================================================
# 4. KEY FEATURES
# ============================================================

add_heading_styled(doc, "4. Key Features", level=1)

# --- 4.1 ---
add_heading_styled(doc, "4.1  Interactive Charging Map", level=2)

add_body(doc,
    "The core of WattsOn is a premium, dark-themed Google Maps interface displaying every "
    "verified charging station in Egypt with custom markers and rich information cards.")

features_41 = [
    "435+ verified stations with verification badges and provider logos",
    "Station detail cards with connector types, power output (kW), pricing, and community ratings",
    "Advanced filters: connector type (CCS2, Type 2, CHAdeMO), power level, provider name",
    "One-tap navigation to any station via Google Maps",
    "Desktop: split-pane layout with scrollable station list alongside the map",
    "Mobile: full-screen map with swipeable bottom sheet station cards",
    "Cluster markers that expand at higher zoom levels for dense areas",
]
for f in features_41:
    add_bullet(doc, f)

# --- 4.2 ---
add_heading_styled(doc, "4.2  Community-Driven Station Status", level=2)

add_body(doc,
    "WattsOn is the first platform in Egypt to provide real-time, crowdsourced availability "
    "data for charging stations -- verified by proximity.",
    bold_phrases=["first platform in Egypt"])

features_42 = [
    "Real-time status reporting: Available, Some Free, Busy, Out of Service",
    "Proximity-locked reporting: only users within 100 meters can submit a status update, preventing spam and ensuring accuracy",
    "Status timestamps showing when each report was submitted and how recent it is",
    "Confidence indicators: Recent (green), A while ago (yellow), Old (grey)",
    "Community reliability score derived from aggregated star ratings",
    "Station reviews with star ratings and quick tags (clean, fast, reliable)",
]
for f in features_42:
    add_bullet(doc, f)

# --- 4.3 ---
add_heading_styled(doc, "4.3  AI Copilot (Claude-Powered)", level=2)

add_body(doc,
    "Every WattsOn user has access to an AI assistant powered by Anthropic's Claude, "
    "trained on real station data and the user's vehicle profile.")

features_43 = [
    "Smart home dashboard with time-aware greetings and proactive charging insights",
    'Natural language queries: "Find a fast charger near Maadi" returns specific station recommendations with distance and power',
    "Real station data integration: the AI knows the user's 15 nearest stations by name, distance, connector types, and power output",
    "Visual response cards: interactive station cards with Navigate buttons, cost comparisons, and battery health summaries",
    "Multi-turn conversation with full context memory across the session",
    "Action execution: the AI can open the trip planner, navigate to a station, or pull up vehicle analytics on command",
]
for f in features_43:
    add_bullet(doc, f)

# --- 4.4 ---
add_heading_styled(doc, "4.4  AI Trip Planner", level=2)

add_body(doc,
    "Plan long-distance trips anywhere in Egypt with intelligent charging stop optimization.")

features_44 = [
    "Route planning between any Egyptian cities (Cairo to Hurghada, Cairo to Alexandria, and more)",
    "Real drive times and distances calculated via the Google Directions API",
    "Battery consumption modeling based on actual vehicle specifications and average speed",
    "Two optimization modes: Quick Stops (2-3 shorter charges) vs Fewer Stops (1-2 longer charges)",
    "Nearby attractions at each charging stop: restaurants, cafes, and scenic spots via Google Places API",
    "Destination autocomplete with 38+ Egyptian cities pre-loaded",
    "Save and revisit trip plans",
]
for f in features_44:
    add_bullet(doc, f)

add_page_break(doc)

# --- 4.5 ---
add_heading_styled(doc, "4.5  Vehicle Dashboard", level=2)

add_body(doc,
    "Deep vehicle intelligence for 140+ EV models across 41 global brands.")

features_45 = [
    "Comprehensive vehicle database: Tesla, BYD, BMW, Mercedes-Benz, Hyundai, Kia, Volkswagen, and 34 more brands",
    "AI Battery Health Report: degradation estimate, charge cycle count, optimal range analysis, and estimated battery life remaining",
    "Egypt heat impact analysis: accounts for 35+ degree Celsius climate factors on battery performance",
    "Consumption analytics: kWh per 100km, cost per kilometer, monthly spending chart, and CO2 emissions saved",
    "Charging pattern insights: preferred charging time, DC vs AC charging ratio, and top visited stations",
    "Personalized AI tips based on vehicle model, driving patterns, and local conditions",
]
for f in features_45:
    add_bullet(doc, f)

# --- 4.6 ---
add_heading_styled(doc, "4.6  EV News Magazine", level=2)

add_body(doc,
    "A built-in news experience that keeps EV drivers informed with curated, positive content.")

features_46 = [
    "Live RSS feeds from 5 top EV publications: Electrek, InsideEVs, CleanTechnica, The Driven, Green Car Reports",
    "Magazine-style layout with hero images, editorial cards, and video overlays",
    "Category filters: Trending, Global, Tech, Market, Reviews",
    "AI-curated article suggestions based on the user's vehicle brand",
    "Daily content refresh with new articles every day",
    "Negative news filtering: recalls, crashes, and lawsuits are automatically excluded",
]
for f in features_46:
    add_bullet(doc, f)

# --- 4.7 ---
add_heading_styled(doc, "4.7  Submit a Station", level=2)

add_body(doc,
    "WattsOn grows through its community. Users can submit new stations they discover in the field.")

features_47 = [
    "Crowdsourced station submissions with a streamlined mobile-first form",
    "GPS auto-fill: the user's current location is pre-populated for accuracy",
    "Provider selection from the full list of known Egyptian providers, with an option to add new ones",
    "Connector type and power level specification",
    "Community verification pipeline: 3 independent confirmations required before a station appears on the map as verified",
    "Verification progress bars visible to all users, encouraging participation",
]
for f in features_47:
    add_bullet(doc, f)

# --- 4.8 ---
add_heading_styled(doc, "4.8  Proximity Intelligence", level=2)

add_body(doc,
    "WattsOn uses the device's GPS to create contextual, location-aware interactions.")

features_48 = [
    "Auto-detect arrival: when a user comes within 100m of a station, a prompt appears to report the station's current status",
    "Auto-detect departure: when the user leaves, a prompt appears to rate the charging experience",
    "48-hour rating window: users can submit a rating within 48 hours of visiting a station",
    "Visit history stored locally on-device for privacy",
]
for f in features_48:
    add_bullet(doc, f)

add_page_break(doc)

# ============================================================
# 5. UNIQUE VALUE PROPOSITION
# ============================================================

add_heading_styled(doc, "5. Unique Value Proposition", level=1)

add_body(doc,
    "WattsOn occupies a unique position in the Egyptian market. No other product combines "
    "aggregation, community data, and AI intelligence into a single platform.")

# UVP comparison table
table = doc.add_table(rows=7, cols=3)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ["Differentiator", "WattsOn", "Competitors"]
for i, h in enumerate(headers):
    cell = table.cell(0, i)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styled_run(p, h, bold=True, size=10, color=WHITE)
    set_cell_shading(cell, "0F172A")

rows_data = [
    ("All providers on one map",     "Yes -- 12+ providers", "Single provider only"),
    ("Real-time crowdsourced status", "Yes -- proximity-verified", "Not available"),
    ("AI copilot",                    "Yes -- Claude-powered", "Not available"),
    ("Intelligent trip planning",     "Yes -- with charging optimization", "Basic or none"),
    ("Vehicle battery analytics",     "Yes -- AI health reports", "Not available"),
    ("Premium UX / dark theme",       "Yes -- Tesla/Rivian-grade", "Basic utility apps"),
]

for r, (diff, wattson, comp) in enumerate(rows_data, start=1):
    c0 = table.cell(r, 0)
    styled_run(c0.paragraphs[0], diff, bold=True, size=10)
    c1 = table.cell(r, 1)
    styled_run(c1.paragraphs[0], wattson, size=10, color=BRAND_GREEN)
    c2 = table.cell(r, 2)
    styled_run(c2.paragraphs[0], comp, size=10, color=BRAND_LIGHT)

doc.add_paragraph()  # spacer

uvp_bullets = [
    ("Only aggregator in Egypt: ", "No other app shows ALL providers on one map."),
    ("Community-driven availability: ", "First in Egypt to offer real-time crowdsourced station status with proximity verification."),
    ("AI-first approach: ", "Every feature enhanced by Claude AI -- not just a map, but an intelligent copilot for the EV driver."),
    ("Proximity-verified data: ", "Reports locked to 100m range ensure trustworthy, spam-free data."),
    ("Premium UX: ", "Dark theme, Space Grotesk typography, gradient navigation tabs, and glowing animations deliver a Tesla/Rivian-level experience."),
    ("435+ stations: ", "The largest verified station database in Egypt."),
]
for prefix, text in uvp_bullets:
    add_bullet(doc, text, bold_prefix=prefix)

add_page_break(doc)

# ============================================================
# 6. TECHNOLOGY STACK
# ============================================================

add_heading_styled(doc, "6. Technology Stack", level=1)

add_body(doc,
    "WattsOn is built on a modern, scalable technology stack chosen for rapid iteration, "
    "cross-platform reach, and real-time capabilities.")

tech_table = doc.add_table(rows=9, cols=3)
tech_table.style = "Table Grid"
tech_table.alignment = WD_TABLE_ALIGNMENT.CENTER

tech_headers = ["Layer", "Technology", "Purpose"]
for i, h in enumerate(tech_headers):
    cell = tech_table.cell(0, i)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styled_run(p, h, bold=True, size=10, color=WHITE)
    set_cell_shading(cell, "0F172A")

tech_rows = [
    ("Frontend",          "React Native + Expo",                  "Cross-platform: iOS, Android, and Web from a single codebase"),
    ("Backend",           "Supabase",                             "PostgreSQL database, authentication, Edge Functions, and real-time subscriptions"),
    ("AI Engine",         "Anthropic Claude API (claude-sonnet-4-20250514)", "Natural language copilot, battery analysis, trip optimization"),
    ("Maps",              "Google Maps JavaScript API",           "Map rendering, Directions, Places, and Geocoding services"),
    ("News",              "Live RSS Feeds + CORS Proxy",          "Real-time articles from 5 top EV publications"),
    ("Hosting",           "Vercel (Web) + Expo (Native)",         "Global CDN for web, managed native builds"),
    ("State Management",  "Zustand + TanStack React Query",       "Lightweight global state and server-state caching"),
    ("Typography",        "Space Grotesk (Google Fonts)",         "Modern geometric sans-serif for premium feel"),
]

for r, (layer, tech, purpose) in enumerate(tech_rows, start=1):
    styled_run(tech_table.cell(r, 0).paragraphs[0], layer, bold=True, size=10)
    styled_run(tech_table.cell(r, 1).paragraphs[0], tech, size=10, color=BRAND_GREEN)
    styled_run(tech_table.cell(r, 2).paragraphs[0], purpose, size=10)

add_page_break(doc)

# ============================================================
# 7. MARKET OPPORTUNITY
# ============================================================

add_heading_styled(doc, "7. Market Opportunity", level=1)

add_body(doc,
    "Egypt's electric vehicle market is entering a period of rapid growth, driven by "
    "government policy, international manufacturer entry, and rising consumer demand.")

market_points = [
    ("40%+ annual growth: ", "Egypt's EV market is expanding at over 40% year-over-year, one of the fastest rates in the Middle East and Africa."),
    ("Government incentives: ", "Customs exemptions on imported EVs, subsidized charging infrastructure rollout, and favorable energy pricing for EV charging."),
    ("100,000 EV target by 2030: ", "The Egyptian government has set an explicit target of 100,000 electric vehicles on the road by 2030."),
    ("Global brands entering: ", "BYD, MG, Tesla, Hyundai, BMW, Mercedes-Benz, and Volkswagen are all establishing or expanding their presence in Egypt."),
    ("$500M infrastructure investment: ", "Charging infrastructure investment is expected to reach $500 million by 2030, creating a massive ecosystem opportunity."),
    ("Platform opportunity: ", "WattsOn is positioned as the platform layer atop this growing infrastructure -- generating revenue through provider partnerships, targeted advertising, data insights, and premium subscription features."),
]

for prefix, text in market_points:
    add_bullet(doc, text, bold_prefix=prefix)

doc.add_paragraph()

add_body(doc,
    "As the only aggregator in the market, WattsOn has the opportunity to become the "
    "default interface between EV drivers and the charging infrastructure -- a position "
    "analogous to Waze for navigation or Fuelio for fuel stations, but with AI intelligence "
    "at its core.",
    bold_phrases=["default interface between EV drivers and the charging infrastructure"])

add_page_break(doc)

# ============================================================
# 8. ROADMAP
# ============================================================

add_heading_styled(doc, "8. Roadmap", level=1)

add_body(doc,
    "WattsOn follows a phased approach to growth, starting with product-market fit in Egypt "
    "and expanding regionally.")

roadmap_table = doc.add_table(rows=5, cols=3)
roadmap_table.style = "Table Grid"
roadmap_table.alignment = WD_TABLE_ALIGNMENT.CENTER

rm_headers = ["Timeline", "Milestone", "Key Deliverables"]
for i, h in enumerate(rm_headers):
    cell = roadmap_table.cell(0, i)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styled_run(p, h, bold=True, size=10, color=WHITE)
    set_cell_shading(cell, "0F172A")

rm_rows = [
    ("Q2 2026", "Public Launch",
     "500+ stations mapped, native iOS and Android apps on App Store and Google Play, marketing launch campaign"),
    ("Q3 2026", "Provider Partnerships",
     "Live availability feeds from partner providers, integrated payment processing, station reservation system"),
    ("Q4 2026", "Enterprise Expansion",
     "Fleet management portal for corporate EV fleets, public API for third-party integrations, B2B analytics dashboard"),
    ("Q1 2027", "Regional Expansion",
     "Launch in UAE and Saudi Arabia, 1,000+ stations across three countries, multi-language support (Arabic, English)"),
]

for r, (timeline, milestone, deliverables) in enumerate(rm_rows, start=1):
    styled_run(roadmap_table.cell(r, 0).paragraphs[0], timeline, bold=True, size=10, color=BRAND_GREEN)
    styled_run(roadmap_table.cell(r, 1).paragraphs[0], milestone, bold=True, size=10)
    styled_run(roadmap_table.cell(r, 2).paragraphs[0], deliverables, size=10)

add_page_break(doc)

# ============================================================
# 9. TEAM & CONTACT
# ============================================================

add_heading_styled(doc, "9. Team & Contact", level=1)

add_body(doc,
    "WattsOn is built by a team passionate about sustainable transportation and technology "
    "innovation in Egypt and the Middle East.")

doc.add_paragraph()

add_body(doc, "Team details to be announced.", size=11)

doc.add_paragraph()

contact_table = doc.add_table(rows=2, cols=2)
contact_table.style = "Table Grid"
contact_table.alignment = WD_TABLE_ALIGNMENT.LEFT

styled_run(contact_table.cell(0, 0).paragraphs[0], "Website", bold=True, size=11)
styled_run(contact_table.cell(0, 1).paragraphs[0], "wattson-ev.vercel.app", size=11, color=BRAND_GREEN)
styled_run(contact_table.cell(1, 0).paragraphs[0], "Contact", bold=True, size=11)
styled_run(contact_table.cell(1, 1).paragraphs[0], "sambawy@gmail.com", size=11, color=BRAND_GREEN)

# Spacer
for _ in range(6):
    doc.add_paragraph()

# Closing
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
styled_run(p, "WattsOn", bold=True, size=16, color=BRAND_GREEN)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
styled_run(p, "Find. Charge. Drive Smart.", italic=True, size=12, color=BRAND_LIGHT)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
styled_run(p, "CONFIDENTIAL -- March 2026", size=9, color=RGBColor(0xCC, 0x00, 0x00))

# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------

output_dir = "/Users/bistrocloud/Documents/EV Charging Aggregator/docs"
os.makedirs(output_dir, exist_ok=True)
output_path = os.path.join(output_dir, "WattsOn_Concept_Document.docx")
doc.save(output_path)
print(f"Document saved to: {output_path}")
