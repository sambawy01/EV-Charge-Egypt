#!/usr/bin/env python3
"""Generate WattsOn vs Zap-Map vs PlugShare Feature Comparison Word Document."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x0A, 0x0E, 0x1A)
    hs.font.name = 'Calibri'

doc.styles['Heading 1'].font.size = Pt(22)
doc.styles['Heading 2'].font.size = Pt(16)
doc.styles['Heading 3'].font.size = Pt(13)

CYAN = RGBColor(0x00, 0xD4, 0xFF)
DARK_BG = RGBColor(0x0A, 0x0E, 0x1A)
GREEN_HEX = "00AA55"
RED_HEX = "CC3333"
CYAN_HEX = "00D4FF"
GRAY_HEX = "999999"
YES_COLOR = RGBColor(0x00, 0xAA, 0x55)
NO_COLOR = RGBColor(0xCC, 0x33, 0x33)
PARTIAL_COLOR = RGBColor(0xDD, 0x99, 0x00)

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def make_comparison_table(doc, rows_data):
    """Create a 4-column comparison table: Feature | WattsOn | Zap-Map | PlugShare"""
    table = doc.add_table(rows=1 + len(rows_data), cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    headers = ["Feature", "WattsOn", "Zap-Map", "PlugShare"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)
        set_cell_shading(cell, "0A0E1A")

    # Data
    for r, row in enumerate(rows_data):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9.5)
                    if c > 0:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # Color code yes/no
                    v = str(val).strip().upper()
                    if v in ("YES", "BEST"):
                        run.font.color.rgb = YES_COLOR
                        run.bold = True
                    elif v in ("NO", "NONE"):
                        run.font.color.rgb = NO_COLOR
                    elif v in ("PARTIAL", "LIMITED", "BASIC"):
                        run.font.color.rgb = PARTIAL_COLOR
            if r % 2 == 0:
                set_cell_shading(cell, "F0F4F8")

    return table

def make_detail_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)
        set_cell_shading(cell, "0A0E1A")

    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9.5)
            if r % 2 == 0:
                set_cell_shading(cell, "F0F4F8")
    return table

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)

# ══════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════
for _ in range(5):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("WattsOn vs Zap-Map vs PlugShare")
run.font.size = Pt(36)
run.bold = True
run.font.color.rgb = CYAN

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Head-to-Head Feature Comparison")
run.font.size = Pt(22)
run.font.color.rgb = DARK_BG

doc.add_paragraph()
doc.add_paragraph()

tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tagline.add_run("How Egypt's first AI-powered EV charging aggregator\nstacks up against the global leaders")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.italic = True

doc.add_paragraph()
doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run(f"Prepared: {datetime.date.today().strftime('%B %Y')}\nConfidential")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# AT A GLANCE
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('Company Overview', level=1)

make_detail_table(doc,
    ["", "WattsOn", "Zap-Map", "PlugShare"],
    [
        ["Founded", "2024", "2014", "2009"],
        ["HQ", "El Gouna, Egypt", "Bristol, UK", "Los Angeles, USA"],
        ["Owner", "Independent (founder-led)", "Independent (VC-backed)", "EVgo Inc. (acquired $25M)"],
        ["Funding", "Bootstrapped", "GBP 9M Series A (2022)", "$25M acquisition (2021)"],
        ["Valuation", "Pre-revenue", "GBP 26.3M", "Part of EVgo ($1.5B+)"],
        ["Employees", "1 (founder)", "~92", "Part of EVgo (~900)"],
        ["Market", "Egypt (MENA planned)", "UK + 5 EU countries", "200+ countries globally"],
        ["Stations Mapped", "435 (12 providers)", "118,321 (45,561 locations)", "750,000+"],
        ["Users", "Early stage", "2M+ downloads, 1M+ active", "3.5M+ registered"],
        ["App Rating", "Web app (no store yet)", "4.5-4.6/5", "4.7/5 (135K+ ratings)"],
        ["Model", "Aggregator + AI", "Aggregator + Payment", "Community + Data"],
        ["Revenue", "Pre-revenue", "~$6M/yr (est.)", "Part of EVgo revenue"],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# MASTER COMPARISON TABLE
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('Master Feature Comparison', level=1)

doc.add_paragraph(
    "The following table compares every major feature across all three platforms. "
    "Green = fully implemented, Red = not available, Yellow = partial/limited implementation."
)
doc.add_paragraph()

make_comparison_table(doc, [
    # --- MAP & DISCOVERY ---
    ["MAP & DISCOVERY", "", "", ""],
    ["Interactive map with markers", "Yes", "Yes", "Yes"],
    ["Station count", "435", "118,321", "750,000+"],
    ["Multi-provider aggregation", "Yes (12 providers)", "Yes (40+ networks)", "Yes (50+ networks)"],
    ["Search by location/address", "Yes", "Yes", "Yes"],
    ["Map + list view", "Yes (split-pane desktop)", "Yes", "Yes"],
    ["Color-coded status markers", "Yes (4 states)", "Yes", "Yes"],
    ["Dark/light map theme", "Yes (dark default)", "Yes", "No"],
    ["Desktop responsive layout", "Yes (split-pane >768px)", "Yes (web app)", "Yes (web app)"],
    # --- FILTERS ---
    ["FILTERS", "", "", ""],
    ["Connector type filter", "Yes (CCS, CHAdeMO, Type 2, GBT)", "Yes", "Yes"],
    ["Charging speed filter", "Yes (22/50/100+ kW)", "Yes (4 tiers)", "Yes (L1/L2/DCFC)"],
    ["Network/provider filter", "Yes (12 providers)", "Yes", "Yes (20+ networks)"],
    ["Price filter (max kWh)", "Yes", "Yes", "Partial (free/paid)"],
    ["Amenities filter", "Yes (wifi, restaurant, bathroom, mall, shade)", "No", "Yes (food, lodging, etc.)"],
    ["Vehicle compatibility filter", "Yes (auto from profile)", "Yes", "Yes (auto from profile)"],
    ["Availability filter", "Yes", "Yes", "Yes"],
    ["Saved filter presets", "No", "Yes (Premium)", "No"],
    # --- STATION DATA ---
    ["STATION DATA", "", "", ""],
    ["Provider/network name", "Yes", "Yes", "Yes"],
    ["Connector types & count", "Yes", "Yes", "Yes"],
    ["Power rating (kW)", "Yes", "Yes", "Yes"],
    ["Pricing information", "Yes", "Yes (live PAYG)", "Yes"],
    ["Live availability status", "Yes (community + realtime)", "Yes (80%+ via CPO APIs)", "Partial (hybrid)"],
    ["User ratings & reviews", "Yes (star rating)", "Yes (star rating)", "Yes (PlugScore 1-10)"],
    ["Community comments/tips", "Yes (status reports)", "Yes (Zap-Chat)", "Yes (reviews + tips)"],
    ["Photos", "No", "Yes (community)", "Yes (725K+ photos)"],
    ["Amenities nearby", "Yes", "No", "Yes"],
    ["Hours of operation", "No", "Yes", "Yes"],
    # --- COMMUNITY ---
    ["COMMUNITY FEATURES", "", "", ""],
    ["Community status reporting", "Yes", "Yes (check-ins)", "Yes (check-ins)"],
    ["Proximity verification (GPS)", "Yes (100m lock)", "No (honor system)", "No (honor system)"],
    ["Rate limiting", "Yes (5/min)", "No", "No"],
    ["User reviews & ratings", "Yes (48h post-visit)", "Yes", "Yes (5.8M+ reviews)"],
    ["Photo uploads", "No", "Yes", "Yes (725K+)"],
    ["Station submission (crowdsource)", "Yes (3-verification)", "Yes (moderated)", "Yes"],
    ["Home charger sharing", "No", "No", "Yes"],
    ["User profiles", "Yes", "Yes", "Yes (badges, messaging)"],
    ["In-app user messaging", "No", "No", "Yes"],
    ["Community forum", "No", "Yes (Zap-Map Cafe)", "No"],
    ["Gamification / badges", "No", "No", "Yes"],
    # --- AI & INTELLIGENCE ---
    ["AI & SMART FEATURES", "", "", ""],
    ["AI copilot / assistant", "Yes (Claude API)", "No", "No"],
    ["AI battery health analysis", "Yes", "No", "No"],
    ["AI trip optimization", "Yes", "No", "No"],
    ["AI charging recommendations", "Yes (context-aware)", "No", "No"],
    ["Smart vehicle-based filtering", "Yes", "Yes", "Yes"],
    ["Predictive availability", "No", "No", "No"],
    # --- TRIP PLANNING ---
    ["TRIP PLANNING", "", "", ""],
    ["Route planner", "Yes (corridor search)", "Yes (3 modes)", "Yes (basic)"],
    ["Vehicle-specific range calc", "Yes (332 vehicles)", "Yes", "Yes"],
    ["Charging strategy options", "Yes (quick vs fewer stops)", "Yes (auto/suggest/manual)", "No"],
    ["Waypoint support", "No", "Yes", "Yes"],
    ["Elevation/terrain modeling", "No", "Yes", "Yes (visual)"],
    ["Weather/speed adjustment", "No", "Yes (speed/incline)", "No"],
    ["Real-time battery input", "No", "Yes", "No"],
    ["Save route plans", "No", "Yes (3 free, unlimited premium)", "No"],
    ["Popular route presets", "Yes (6 Egypt routes)", "No", "No"],
    ["Egypt route coverage", "Best (435 local stations)", "None", "Minimal"],
    # --- PAYMENT ---
    ["PAYMENT & CHARGING", "", "", ""],
    ["In-app payment", "Yes (wallet system)", "Yes (Zap-Pay, 40+ networks)", "Limited (EVgo CA only)"],
    ["Start/stop charging", "No", "Yes (via Zap-Pay)", "Limited"],
    ["Cross-network payment", "No", "Yes (80K+ chargers)", "No"],
    ["Physical RFID card", "No", "Yes (free w/ Premium)", "No"],
    ["Mobile money (Egypt)", "Yes (Vodafone/Orange/Etisalat Cash)", "No", "No"],
    ["Fawry integration", "Yes", "No", "No"],
    ["Auto top-up wallet", "Yes", "No", "No"],
    ["Transaction history", "Yes", "Yes", "Yes"],
    ["5% charging discount", "No", "Yes (Premium, 50kWh/mo cap)", "No"],
    # --- VEHICLE ---
    ["VEHICLE FEATURES", "", "", ""],
    ["Vehicle database", "332 models, 88 brands", "EV models (gaps reported)", "Large (US/EU focused)"],
    ["Vehicle dashboard", "Yes (full screen)", "No", "No"],
    ["Battery health scoring", "Yes (AI-powered)", "No", "No"],
    ["Consumption analytics", "Yes (AI estimates)", "No", "No"],
    ["Multi-vehicle profiles", "Yes", "Yes (Premium)", "Yes"],
    ["Charging stats tracking", "Yes (sessions, kWh, spend, CO2)", "No", "Partial (check-in count)"],
    # --- CONTENT ---
    ["CONTENT & NEWS", "", "", ""],
    ["EV news magazine", "Yes (live RSS, 5 sources)", "No", "No"],
    ["Editorial guides", "No", "Yes (web, beginner + advanced)", "No"],
    ["Charging statistics page", "No", "Yes (UK infrastructure data)", "No"],
    ["Cost calculators", "No", "Yes (journey, home, public)", "No"],
    # --- PLATFORM ---
    ["PLATFORM & INTEGRATION", "", "", ""],
    ["iOS app", "No (web app)", "Yes", "Yes"],
    ["Android app", "No (web app)", "Yes", "Yes"],
    ["Web app", "Yes (primary)", "Yes", "Yes"],
    ["Apple CarPlay", "No", "Yes (Premium)", "Yes"],
    ["Android Auto", "No", "Yes (Premium)", "No"],
    ["Apple Watch", "No", "No", "Yes"],
    ["Apple Vision Pro", "No", "No", "Yes"],
    ["Push notifications", "Yes (Expo)", "Yes", "Yes"],
    # --- LOCALIZATION ---
    ["LOCALIZATION", "", "", ""],
    ["Languages supported", "2 (Arabic + English)", "1 (English only)", "29 languages"],
    ["Arabic support", "Yes (229 keys, RTL)", "No", "No"],
    ["RTL layout support", "Yes", "No", "No"],
    # --- SUBSCRIPTION ---
    ["PRICING", "", "", ""],
    ["Free tier", "Yes (full access)", "Yes (core features)", "Yes (full access)"],
    ["Premium subscription", "Not yet", "GBP 29.99/yr", "$14.99/yr (ad-free only)"],
    ["Ad-free option", "N/A (no ads yet)", "Yes (Premium)", "Yes ($14.99/yr)"],
    # --- BUSINESS ---
    ["B2B & DATA", "", "", ""],
    ["Fleet management", "Planned", "Yes (Zap-Pay for Fleets)", "No"],
    ["Data licensing / API", "No", "Yes (Spark APIs, Insights)", "Yes (DataTool, API)"],
    ["CPO dashboard", "No", "No", "Yes"],
    ["Research panel", "No", "No", "Yes (63K+ drivers)"],
    ["Advertising platform", "No", "Yes", "Yes"],
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# DEEP DIVES
# ══════════════════════════════════════════════════════════════════════════

doc.add_heading('Deep Dive: Map & Station Discovery', level=1)

doc.add_paragraph(
    "All three apps center on an interactive charging station map, but they differ significantly "
    "in scale, data sources, and presentation."
)

make_detail_table(doc,
    ["Dimension", "WattsOn", "Zap-Map", "PlugShare"],
    [
        ["Total Stations", "435 verified (Egypt)", "118,321 chargers / 45,561 locations (UK + EU)", "750,000+ (200+ countries)"],
        ["Data Source", "Supabase DB (curated) + community submissions", "CPO API feeds + community check-ins", "Network APIs + massive crowd-sourcing"],
        ["Coverage Depth", "Deep in Egypt (16 governorates, 12 providers)", "Deep in UK (~85% coverage), expanding EU", "Wide but shallow per-country"],
        ["Live Status", "Supabase Realtime + community reports", "80%+ via direct CPO APIs", "Hybrid: some network APIs + check-ins"],
        ["Map Theme", "Dark mode default (futuristic)", "Dark/light toggle", "Standard (no dark mode)"],
        ["Desktop Layout", "Split-pane (list + map side-by-side)", "Standard web map", "Standard web map"],
        ["Mobile Layout", "Full-screen map + bottom sheet overlays", "Full map + list", "Full map + list"],
        ["Station Status States", "4 (available, partial, occupied, offline)", "4 (available, charging, offline, unknown)", "3+ (open, under repair, coming soon)"],
    ]
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("WattsOn's advantage: ")
run.bold = True
p.add_run(
    "In Egypt, WattsOn has the deepest station database (435 across 12 providers vs. minimal coverage "
    "on PlugShare and zero on Zap-Map). The split-pane desktop layout and dark-mode-first design "
    "offer a more modern UX than either competitor."
)

doc.add_page_break()

# --- Community Deep Dive ---
doc.add_heading('Deep Dive: Community & Data Quality', level=1)

doc.add_paragraph(
    "Community features are a critical differentiator for charging aggregators. The quality of "
    "crowd-sourced data directly impacts user trust and platform stickiness."
)

make_detail_table(doc,
    ["Dimension", "WattsOn", "Zap-Map", "PlugShare"],
    [
        ["Reporting Method", "Status report popup (4 options)", "Check-in with outcome + text + photo", "Check-in with outcome + text + photo"],
        ["Proximity Verification", "GPS-locked at 100m radius", "No verification (honor system)", "No verification (honor system)"],
        ["Rate Limiting", "5 reports/minute (client-side)", "None documented", "None documented"],
        ["Data Trust Level", "High (physically verified presence)", "Medium (unverified location)", "Medium (unverified location)"],
        ["Review Scale", "5-star (48h post-visit window)", "Star rating", "PlugScore 1-10 (algorithmic)"],
        ["Photo Support", "Not yet", "Yes (community uploads)", "Yes (725,000+ photos)"],
        ["Community Scale", "Early stage", "1M+ active users", "3.5M+ registered, 5.8M+ reviews"],
        ["Station Submission", "Yes (3 community verifications required)", "Yes (moderated by data team)", "Yes (open submission)"],
        ["ICE'd Reporting", "Not yet", "Yes", "Yes (reason codes)"],
        ["User Messaging", "No", "No", "Yes (in-app DM)"],
    ]
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("WattsOn's advantage: ")
run.bold = True
p.add_run(
    "The 100-meter GPS proximity lock is a genuine innovation. Neither PlugShare nor Zap-Map verify "
    "that reporters are physically at the station, which means anyone can submit false status reports. "
    "WattsOn's approach guarantees data integrity — every report comes from someone who is actually there. "
    "The 3-verification system for new station submissions adds another trust layer that competitors lack."
)

doc.add_page_break()

# --- AI Deep Dive ---
doc.add_heading('Deep Dive: AI & Intelligence', level=1)

doc.add_paragraph(
    "This is WattsOn's most significant differentiator. No EV charging app in the world — "
    "including PlugShare, Zap-Map, ChargePoint, or any other — offers a conversational AI assistant."
)

make_detail_table(doc,
    ["AI Capability", "WattsOn", "Zap-Map", "PlugShare"],
    [
        ["AI Copilot", "Yes (Claude API via Supabase Edge Function)", "No", "No"],
        ["Conversational interface", "Yes (natural language chat)", "No", "No"],
        ["Station recommendations", "Yes (context-aware, uses real station data)", "No", "No"],
        ["Battery health analysis", "Yes (AI degradation estimates, health scoring)", "No", "No"],
        ["Charging optimization", "Yes (cost, speed, convenience trade-offs)", "No", "No"],
        ["Trip planning assistance", "Yes (AI-optimized stops)", "Algorithmic only", "Basic algorithm"],
        ["Egypt-specific knowledge", "Yes (providers, routes, climate, pricing)", "N/A", "N/A"],
        ["Context passed to AI", "Vehicle specs, battery info, 15 nearest stations, user location", "N/A", "N/A"],
        ["Fallback system", "Keyword-based responses if API unavailable", "N/A", "N/A"],
    ]
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Why this matters: ")
run.bold = True
p.add_run(
    "EV charging is complex — connector compatibility, charging speeds, pricing structures, battery management, "
    "and route optimization all require knowledge that most drivers don't have. WattsOn's AI copilot makes this "
    "accessible through natural conversation. A driver can ask 'Where should I charge on the way to Hurghada?' "
    "and get a personalized answer based on their vehicle, battery state, and the actual stations along the route. "
    "No other app in the world can do this."
)

doc.add_page_break()

# --- Trip Planning Deep Dive ---
doc.add_heading('Deep Dive: Trip Planning', level=1)

make_detail_table(doc,
    ["Capability", "WattsOn", "Zap-Map", "PlugShare"],
    [
        ["Planning Approach", "Corridor search along route + AI optimization", "3 modes: Auto / Suggest / Manual", "Basic start-end with stops"],
        ["Vehicle Range Calc", "Yes (332 models)", "Yes (EV database)", "Yes (EV database)"],
        ["Charging Strategy", "Quick charging vs. fewer stops", "Auto-selects optimal stops", "No strategy options"],
        ["Elevation Modeling", "No", "Yes (reduces range on inclines)", "Yes (visual elevation profile)"],
        ["Speed Adjustment", "No", "Yes (reduces range over 50mph)", "No"],
        ["Weather Factor", "No", "No", "No"],
        ["Waypoint Support", "No", "Yes", "Yes"],
        ["Save Routes", "No", "Yes (3 free / unlimited premium)", "No"],
        ["Popular Presets", "6 Egypt routes (Hurghada, Alex, Sharm, Ain Sokhna, El Alamein, El Gouna)", "No", "No"],
        ["Geographic Coverage", "Egypt (435 stations)", "UK + Ireland only", "Global (but shallow per-country)"],
        ["In-car Handoff", "No", "Yes (CarPlay/AA + Apple/Google Maps)", "No"],
        ["AI Integration", "Yes (Claude optimizes stops)", "No", "No"],
    ]
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("WattsOn's advantage: ")
run.bold = True
p.add_run(
    "For Egyptian routes, WattsOn is the only viable trip planner. Zap-Map's planner is UK/Ireland only, "
    "and PlugShare has minimal Egypt data. WattsOn's corridor search finds stations along the actual route, "
    "and the AI integration can optimize stops based on the driver's specific vehicle and preferences. "
    "Zap-Map has more sophisticated range modeling (elevation, speed), which is an area for WattsOn to improve."
)

doc.add_page_break()

# --- Payment Deep Dive ---
doc.add_heading('Deep Dive: Payment & Charging', level=1)

make_detail_table(doc,
    ["Capability", "WattsOn", "Zap-Map", "PlugShare"],
    [
        ["Payment System", "Wallet-based (top-up)", "Zap-Pay (direct charge)", "Pay with PlugShare (limited)"],
        ["Payment Methods", "Card, Fawry, InstaPay, Vodafone Cash, Orange Cash, Etisalat Cash", "Card, Apple Pay, Google Pay", "Credit card"],
        ["Network Coverage", "Planned (not live yet)", "40+ networks, 80K+ chargers", "EVgo in California only"],
        ["Physical Card", "No", "RFID card (free w/ Premium)", "No"],
        ["Start/Stop Charging", "No", "Yes (app or card tap)", "Limited"],
        ["Real-time Session Monitor", "No", "Yes", "Yes"],
        ["Charging Discount", "No", "5% via Premium (50kWh/mo cap)", "No"],
        ["Mobile Money", "Yes (3 Egypt mobile wallets)", "No", "No"],
        ["Local Payment (Fawry)", "Yes", "No", "No"],
        ["Auto Top-up", "Yes (threshold-based)", "No", "No"],
    ]
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Analysis: ")
run.bold = True
p.add_run(
    "Zap-Pay is the gold standard for cross-network EV payment, covering 80,000+ chargers across 40+ networks "
    "with a physical RFID card option. PlugShare's payment is extremely limited (EVgo California only). "
    "WattsOn has built a wallet system with Egypt-specific payment methods (Fawry, mobile money) that neither "
    "competitor supports — critical for the Egyptian market where credit card penetration is low. "
    "The gap is in CPO integration for session start/stop, which is WattsOn's next priority."
)

doc.add_page_break()

# --- Vehicle Deep Dive ---
doc.add_heading('Deep Dive: Vehicle Features', level=1)

make_detail_table(doc,
    ["Capability", "WattsOn", "Zap-Map", "PlugShare"],
    [
        ["Vehicle Database Size", "332 models from 88 brands", "Unknown (some gaps reported)", "Large (US/EU focus)"],
        ["Chinese EV Coverage", "Extensive (BYD, Chery, Geely, Changan, GAC, Jetour, MG, etc.)", "Limited", "Limited"],
        ["Dedicated Dashboard", "Yes (full-screen vehicle view)", "No", "No"],
        ["Battery Capacity Display", "Yes (kWh)", "Used for range calc only", "Used for range calc only"],
        ["AI Battery Health Score", "Yes (degradation estimate, health grade)", "No", "No"],
        ["AI Optimization Tips", "Yes (temperature, charging patterns)", "No", "No"],
        ["Charging Stats", "Sessions, kWh, spend (EGP), CO2 saved", "No personal stats", "Check-in count only"],
        ["Multi-vehicle Support", "Yes", "Yes (Premium only)", "Yes"],
    ]
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("WattsOn's advantage: ")
run.bold = True
p.add_run(
    "WattsOn's vehicle dashboard is a unique feature category. No competitor offers a dedicated vehicle view "
    "with AI-powered battery health analysis. The extensive Chinese EV coverage (BYD, Chery, Geely, etc.) "
    "is critical for Egypt where Chinese brands dominate new EV sales. PlugShare and Zap-Map's databases "
    "are US/EU-centric with limited support for Chinese models popular in MENA markets."
)

doc.add_page_break()

# --- Localization Deep Dive ---
doc.add_heading('Deep Dive: Localization & Accessibility', level=1)

make_detail_table(doc,
    ["Dimension", "WattsOn", "Zap-Map", "PlugShare"],
    [
        ["Languages", "2 (Arabic + English)", "1 (English only)", "29 (no Arabic)"],
        ["Arabic Support", "Yes (229 translation keys)", "No", "No"],
        ["RTL Layout", "Yes (full RTL support)", "No", "No"],
        ["Local Currency", "EGP (Egyptian Pound)", "GBP", "Multiple (but not EGP)"],
        ["Local Payment Methods", "Fawry, InstaPay, Vodafone/Orange/Etisalat Cash", "UK cards, Apple/Google Pay", "Credit card only"],
        ["Local Content", "Egypt-specific routes, providers, EV news", "UK-specific guides, statistics", "Global (US-centric content)"],
        ["Regional Knowledge", "Egypt governorates, providers, climate, pricing", "UK charging networks, regulations", "Global but shallow per-region"],
    ]
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("WattsOn's advantage: ")
run.bold = True
p.add_run(
    "Neither Zap-Map (English only) nor PlugShare (29 languages but no Arabic) serve Arabic-speaking users. "
    "For Egypt's 110M+ population where Arabic is the primary language, this is not optional — it's essential. "
    "WattsOn's 229-key Arabic translation with full RTL layout support makes it the only viable option "
    "for the majority of Egyptian EV drivers."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# SCORING SUMMARY
# ══════════════════════════════════════════════════════════════════════════

doc.add_heading('Feature Scoring Summary', level=1)

doc.add_paragraph(
    "Each category is scored 0-10 based on feature completeness, quality, and market relevance. "
    "Scores reflect the current state of each platform."
)

make_detail_table(doc,
    ["Category", "WattsOn", "Zap-Map", "PlugShare", "Notes"],
    [
        ["Map & Discovery", "7/10", "9/10", "9/10", "WattsOn leads in Egypt; others lead globally"],
        ["Filters", "7/10", "9/10", "8/10", "Zap-Map's saved presets and premium filters lead"],
        ["Station Data", "6/10", "9/10", "9/10", "WattsOn lacks photos, hours; others more complete"],
        ["Community", "8/10", "7/10", "9/10", "WattsOn's proximity lock is best-in-class for data quality"],
        ["AI & Intelligence", "10/10", "0/10", "0/10", "WattsOn is the only platform with AI features"],
        ["Trip Planning", "6/10", "8/10", "5/10", "Zap-Map most sophisticated; WattsOn best for Egypt"],
        ["Payment", "5/10", "9/10", "2/10", "Zap-Pay dominates; WattsOn has Egypt payments, no CPO link"],
        ["Vehicle Features", "9/10", "4/10", "3/10", "WattsOn's dashboard + AI health is unique"],
        ["Content & News", "8/10", "7/10", "1/10", "WattsOn's live RSS is engaging; Zap-Map has guides"],
        ["Localization", "9/10", "3/10", "8/10", "WattsOn's Arabic + RTL is unmatched for MENA"],
        ["Platform Coverage", "4/10", "8/10", "9/10", "WattsOn web-only; others have native apps + CarPlay"],
        ["B2B & Data", "1/10", "7/10", "9/10", "PlugShare's DataTool + Research panel lead"],
        ["TOTAL (avg)", "6.7/10", "6.7/10", "6.0/10", "WattsOn matches Zap-Map overall despite being newer"],
    ]
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Key takeaway: ")
run.bold = True
p.add_run(
    "WattsOn scores competitively with platforms that are 10-15 years older and backed by millions "
    "in funding. Its AI features (10/10), vehicle dashboard (9/10), and localization (9/10) are "
    "category-leading. The gaps are in platform distribution (no native apps yet), B2B products, "
    "and payment CPO integration — all addressable with the existing tech stack."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# WHAT WATTSON CAN LEARN
# ══════════════════════════════════════════════════════════════════════════

doc.add_heading('What WattsOn Can Learn From Each Competitor', level=1)

doc.add_heading('From Zap-Map', level=2)

zap_lessons = [
    ("Zap-Pay cross-network payment: ", "The ability to pay at 40+ networks from one app is Zap-Map's killer feature. WattsOn should prioritize CPO payment integration with Infinity and Sha7en."),
    ("CarPlay / Android Auto: ", "In-car integration is a Premium-gated feature that drives subscriptions. With Expo, WattsOn could build a CarPlay extension."),
    ("Route planning sophistication: ", "Speed and elevation adjustments for range calculation improve accuracy. WattsOn should add these to its corridor search algorithm."),
    ("Saved filter presets: ", "Power users want to save their preferred filter combinations. Simple feature, high engagement value."),
    ("Physical RFID card: ", "A branded charging card is both a payment method and a marketing tool. Consider for WattsOn Plus subscribers."),
    ("Editorial content & guides: ", "Zap-Map's beginner/advanced guides build trust with new EV owners. WattsOn could add Egypt-specific charging guides."),
    ("Charging statistics page: ", "Public infrastructure data builds authority and attracts press/government attention."),
]

for bold, text in zap_lessons:
    add_bullet(doc, text, bold)

doc.add_heading('From PlugShare', level=2)

plug_lessons = [
    ("Community scale & engagement: ", "5.8M reviews and 725K photos create an unassailable data moat. WattsOn should aggressively incentivize community contributions early."),
    ("Photo uploads: ", "User-contributed station photos are extremely valuable for navigation and trust. Add photo support to WattsOn's reporting flow."),
    ("Home charger sharing: ", "Peer-to-peer home charging fills infrastructure gaps. Relevant for Egypt where public stations are sparse."),
    ("B2B data products: ", "PlugShare's DataTool and Research panel generate significant revenue. WattsOn's Egypt-specific data will be valuable to CPOs and government planners."),
    ("29 languages: ", "Global reach through localization. As WattsOn expands to MENA, adding Turkish, French (Morocco), and Urdu (Gulf workers) would open markets."),
    ("PlugScore algorithm: ", "A transparent, recency-weighted reliability score builds trust. WattsOn could implement a similar system using its verified proximity reports."),
    ("CPO Dashboard: ", "Giving charging operators a dashboard to monitor their stations on WattsOn creates a B2B relationship that's hard to break."),
    ("In-app messaging: ", "Direct messaging between users enables community coordination, especially for home charger sharing."),
]

for bold, text in plug_lessons:
    add_bullet(doc, text, bold)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# WATTSON'S COMPETITIVE MOATS
# ══════════════════════════════════════════════════════════════════════════

doc.add_heading("WattsOn's Competitive Moats vs. Both Competitors", level=1)

doc.add_paragraph(
    "Despite being younger and smaller, WattsOn has built several defensible advantages "
    "that neither Zap-Map nor PlugShare can easily replicate:"
)

moats = [
    ("1. AI Copilot (Neither Competitor Has This)",
     "WattsOn is the only EV charging app in the world with a conversational AI assistant. "
     "The Claude-powered copilot understands real station data, vehicle specs, and Egyptian geography. "
     "Building this requires deep AI integration expertise that CPO apps won't prioritize, and adding AI "
     "to legacy apps like PlugShare (founded 2009) would require fundamental architecture changes."),

    ("2. GPS-Verified Community Data (Both Competitors Use Honor System)",
     "WattsOn's 100-meter proximity lock guarantees every status report comes from someone physically at the station. "
     "PlugShare and Zap-Map both rely on honor-system reporting where anyone can submit from anywhere. "
     "As WattsOn's community grows, this creates a higher-quality dataset that powers better AI recommendations — "
     "a compounding advantage."),

    ("3. Arabic-First for MENA (Neither Competitor Supports Arabic)",
     "PlugShare supports 29 languages but not Arabic. Zap-Map supports only English. "
     "For Egypt (110M people) and the broader MENA region (400M+ Arabic speakers), "
     "WattsOn is the only option. This isn't a feature — it's a market access requirement."),

    ("4. Egypt's Only Complete Station Map",
     "WattsOn's 435 stations across 12 providers is the most comprehensive EV charging database in Egypt. "
     "PlugShare has minimal Egypt coverage. Zap-Map has zero. A global player entering Egypt would need "
     "months of data collection to match WattsOn's coverage — and they'd still lack the community verification."),

    ("5. Vehicle Dashboard with Chinese EV Focus",
     "332 vehicles including extensive Chinese brand coverage (BYD, Chery, Geely, Changan, GAC, Jetour, MG) "
     "reflects the reality of Egypt's EV market where Chinese imports dominate. PlugShare and Zap-Map's databases "
     "are US/EU-centric. The AI battery health feature adds unique value no competitor offers."),

    ("6. Egypt-Specific Payment Methods",
     "Fawry, InstaPay, Vodafone Cash, Orange Cash, and Etisalat Cash cover the payment methods Egyptians "
     "actually use. In a market with low credit card penetration, this is essential. Neither Zap-Map's UK-focused "
     "Zap-Pay nor PlugShare's California-only payment support any Egyptian payment rail."),
]

for title, desc in moats:
    doc.add_heading(title, level=3)
    doc.add_paragraph(desc)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# CONCLUSION
# ══════════════════════════════════════════════════════════════════════════

doc.add_heading('Conclusion', level=1)

doc.add_paragraph(
    "WattsOn competes favorably with platforms that are a decade older and backed by millions in funding. "
    "While Zap-Map and PlugShare lead in scale, payment integration, and B2B products, WattsOn leads in "
    "three categories that neither can replicate quickly:"
)

conclusions = [
    ("AI Intelligence: ", "The only EV charging app with a conversational AI copilot — a genuine global first."),
    ("Data Trust: ", "GPS-verified proximity reporting produces higher-quality community data than any competitor."),
    ("Market Access: ", "Arabic language, Egyptian payment methods, and local station coverage make WattsOn the only viable option for the region."),
]

for bold, text in conclusions:
    add_bullet(doc, text, bold)

doc.add_paragraph()

doc.add_paragraph(
    "The priority roadmap to close remaining gaps:"
)

priorities = [
    ("Q2-Q3 2026: ", "Photo uploads for station reports, CPO payment integration with Infinity"),
    ("Q3-Q4 2026: ", "Native iOS/Android apps via Expo, CarPlay support"),
    ("Q1 2027: ", "PlugScore-style reliability rating, saved filter presets, elevation-aware trip planning"),
    ("Q2 2027: ", "B2B dashboard for CPOs, data licensing API"),
]

for bold, text in priorities:
    add_bullet(doc, text, bold)

doc.add_paragraph()

final = doc.add_paragraph()
run = final.add_run(
    "WattsOn is not trying to out-PlugShare PlugShare. It is building the Zap-Map of MENA — "
    "a region-dominant aggregator with AI capabilities that no global player offers — and the window "
    "to establish this position is open now."
)
run.italic = True

# ── Save ────────────────────────────────────────────────────────────────
output_path = "/Users/bistrocloud/Documents/EV Charging Aggregator/docs/WattsOn_vs_ZapMap_vs_PlugShare_Comparison.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")
